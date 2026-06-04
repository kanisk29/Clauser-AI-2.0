import io
import docx
from groq import Groq
import re
import json
import os
from dotenv import load_dotenv
from langchain_chroma import Chroma
from langchain_community.embeddings import FakeEmbeddings
load_dotenv()

embeddings_func = FakeEmbeddings(size = 384)
vecdb = Chroma(persist_directory="db/chroma_db",embedding_function=embeddings_func)

def analyze_contract_risk(contract_text, contract_type, industry, persona,custom_knowledge=None):

    """
    Sends the contract text, parameters, and custom RAG knowledge to the LLM.
    """
    if custom_knowledge:
        docs = vecdb.similarity_search(query=f"legal risks in {contract_type} contract for {persona} in {industry}",k = 5)
        retrieved_text = "\n\n".join(doc.page_content for doc in docs)
        if retrieved_text:
            system_prompt = f"""
            You are an expert legal AI assistant. Analyze the provided contract.
            Contract Type: {contract_type}
            Industry: {industry}
            Analyzing from the perspective of the: {persona}

            Relevant external knowledge (company playbook / legal clauses):
            {retrieved_text if retrieved_text else "None provided. Rely on standard legal knowledge."}
            
            Identify key risks based on the {persona}'s perspective, prioritizing the RAG data if provided. 
            For each risk, provide:
            1. The specific clause/issue.
            2. The risk level (High, Medium, Low).
            3. Relevant regulatory/compliance references.
            4. Suggested mitigation.
            
            Return the response EXCLUSIVELY as a JSON array of objects with keys. ONLY RETURN THE JSON ARRAY NOTHING ELSE: 
            "clause", "risk_level", "explanation", "reference", "mitigation".
            """
        else:
            system_prompt = f"""
            You are an expert legal AI assistant. Analyze the provided contract.
            Contract Type: {contract_type}
            Industry: {industry}
            Analyzing from the perspective of the: {persona}
            
            Identify key risks based on the {persona}'s perspective.
            For each risk, provide:
            1. The specific clause/issue.
            2. The risk level (High, Medium, Low).
            3. Relevant regulatory/compliance references.
            4. Suggested mitigation.
            
            Return the response EXCLUSIVELY as a JSON array of objects with keys. ONLY RETURN THE JSON ARRAY NOTHING ELSE: 
            "clause", "risk_level", "explanation", "reference", "mitigation".
            """


    else:
        system_prompt = f"""
        You are an expert legal AI assistant. Analyze the provided contract.
        Contract Type: {contract_type}
        Industry: {industry}
        Analyzing from the perspective of the: {persona}
        
        Identify key risks based on the {persona}'s perspective.
        For each risk, provide:
        1. The specific clause/issue.
        2. The risk level (High, Medium, Low).
        3. Relevant regulatory/compliance references.
        4. Suggested mitigation.
        
        Return the response EXCLUSIVELY as a JSON array of objects with keys. ONLY RETURN THE JSON ARRAY NOTHING ELSE: 
        "clause", "risk_level", "explanation", "reference", "mitigation".
        """

    client = Groq(api_key=os.getenv("GROQ_API_KEY"))
    try:
        response = client.chat.completions.create(
            model='llama-3.3-70b-versatile',
            messages=[{"role":"system","content":system_prompt},
                    {"role":"user","content":contract_text[:15000]}]
        ).choices[0].message.content
    except Exception as e:
        print(e)
        return [] 
    print(response)
    response = re.sub(r"```json", "", response)
    response = re.sub(r"```", "", response)
    response = response.strip()
    response = json.loads(response)
    print(response)
    return response

def generate_revised_docx(original_text, accepted_mitigations):
    """
    Generates a new DOCX file with the applied mitigations.
    """
    doc = docx.Document()
    doc.add_heading('Revised Contract (Mitigated)', 0)
    
    doc.add_paragraph("--- ORIGINAL TEXT WITH APPLIED CHANGES ---")
    doc.add_paragraph(original_text[:500] + "...\n[Original text truncated for preview]")
    
    doc.add_heading('Applied Mitigations', level=1)
    for change in accepted_mitigations:
        doc.add_paragraph(f"Clause Adjusted: {change['clause']}", style='List Bullet')
        doc.add_paragraph(f"New Text: {change['mitigation']}")
        
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer