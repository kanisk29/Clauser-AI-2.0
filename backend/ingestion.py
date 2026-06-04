import PyPDF2
import docx

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from retrieval import embeddings_func


def extract_text_from_file(file_path):

    text = ""

    file_type = file_path.split(".")[-1].lower()

    try:

        if file_type == "pdf":

            with open(file_path, "rb") as file:

                pdf_reader = PyPDF2.PdfReader(file)

                for page in pdf_reader.pages:

                    extracted = page.extract_text()

                    if extracted:

                        text += extracted + "\n"

        elif file_type in ["doc", "docx"]:

            doc = docx.Document(file_path)

            for para in doc.paragraphs:

                text += para.text + "\n"

        else:

            return "Unsupported file format"

    except Exception as e:

        return f"Error extracting text: {str(e)}"

    return text


def file_indexer(file_path, filename):

    text = extract_text_from_file(file_path)

    docs = chunker(text, filename)

    vecdb(docs)

    return len(docs)


def chunker(text, filename):

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100
    )

    chunks = splitter.split_text(text)

    docs = [
        Document(
            page_content=chunk,
            metadata={
                "chunk_id": idx,
                "source": filename
            }
        )
        for idx, chunk in enumerate(chunks, 1)
    ]

    return docs


def vecdb(docs):

    vecstore = Chroma(
        persist_directory="db/chroma_db",
        embedding_function=embeddings_func
    )

    vecstore.add_documents(docs)